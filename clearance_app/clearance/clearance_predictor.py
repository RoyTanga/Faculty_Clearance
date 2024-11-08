import PyPDF2
import docx
import re
import pytesseract
from PIL import Image
import pdf2image
import io
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        # Set Tesseract path for Windows
        if os.name == 'nt':  # Windows
            tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
            else:
                logger.error("Tesseract not found. Please install Tesseract-OCR")
        
        # Enhanced keywords with variations and common misspellings
        self.clearance_keywords = {
            'ACADEMIC': [
                'grades', 'submitted', 'academic', 'requirements', 'course', 
                'completion', 'clearance', 'obligations', 'records',
                'standing', 'performance', 'coursework', 'completed',
                'class', 'attendance', 'semester', 'faculty', 'professor',
                'teaching', 'subjects', 'units', 'courses'
            ],
            'FINANCIAL': [
                'fees', 'paid', 'financial', 'payment', 'balance',
                'accounts', 'settled', 'dues', 'outstanding', 'payments',
                'billing', 'receipt', 'transaction', 'money', 'amount',
                'cashier', 'finance', 'accounting', 'paid in full'
            ],
            'LIBRARY': [
                'library', 'books', 'borrowed', 'returned', 'materials',
                'resources', 'dues', 'fines', 'lending', 'borrowing',
                'circulation', 'reference', 'journals', 'publications',
                'media center', 'reading', 'literature'
            ],
            'RESEARCH': [
                'research', 'thesis', 'dissertation', 'paper', 'study',
                'investigation', 'analysis', 'findings', 'methodology',
                'data', 'results', 'conclusion', 'references', 'bibliography',
                'publication', 'journal', 'research work'
            ],
            'EQUIPMENT': [
                'equipment', 'laboratory', 'tools', 'instruments', 'devices',
                'apparatus', 'materials', 'inventory', 'borrowed', 'returned',
                'lab', 'supplies', 'hardware', 'machinery', 'facilities',
                'resources', 'items'
            ],
            'ADMIN': [
                'administrative', 'office', 'department', 'approval',
                'processing', 'verification', 'compliance', 'requirements',
                'documentation', 'forms', 'papers', 'records', 'files',
                'administration', 'management', 'staff', 'personnel'
            ]
        }

    def extract_text_from_pdf(self, pdf_path):
        """Extract text from PDF using multiple methods."""
        text = ""
        
        # Try PyPDF2 first
        try:
            logger.info(f"Attempting to extract text from PDF using PyPDF2: {pdf_path}")
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {str(e)}")

        # If no text was extracted, try pdf2image + OCR
        if not text.strip():
            try:
                logger.info("No text extracted with PyPDF2, attempting pdf2image + OCR")
                poppler_path = self.get_poppler_path()
                logger.info(f"Using poppler path: {poppler_path}")
                
                # Convert PDF to images
                images = pdf2image.convert_from_path(
                    pdf_path,
                    poppler_path=poppler_path
                )
                
                logger.info(f"Successfully converted PDF to {len(images)} images")
                
                # Extract text from each image
                for i, image in enumerate(images):
                    page_text = pytesseract.image_to_string(image)
                    text += page_text + "\n"
                    logger.info(f"Extracted {len(page_text)} characters from page {i+1}")
                    
            except Exception as e:
                logger.error(f"PDF to image conversion failed: {str(e)}")
                
                # If both methods fail, try one last time with just the first page
                try:
                    logger.info("Attempting final extraction with PyPDF2 (first page only)")
                    with open(pdf_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        if len(pdf_reader.pages) > 0:
                            text = pdf_reader.pages[0].extract_text() or ""
                except Exception as e:
                    logger.error(f"Final PDF extraction attempt failed: {str(e)}")

        logger.info(f"Total extracted text length: {len(text)} characters")
        return text

    def get_poppler_path(self):
        """Get the poppler path based on the operating system."""
        if os.name == 'nt':  # Windows
            # Your specific poppler path
            poppler_path = r'C:\Users\lenovo\Downloads\Release-24.08.0-0\poppler-24.08.0\Library\bin'
            
            if os.path.exists(poppler_path):
                return poppler_path
            
            # Fallback paths
            possible_paths = [
                r'C:\Program Files\poppler-24.08.0\Library\bin',
                r'C:\Program Files\poppler\bin',
                r'C:\poppler\bin',
                os.path.join(os.path.dirname(__file__), 'poppler', 'bin'),
            ]
            
            for path in possible_paths:
                if os.path.exists(path):
                    return path
            
            # If no path is found, log an error
            logger.error("Poppler path not found. Please install poppler and add to PATH")
            return None
        else:
            return None  # Return None for Unix-like systems as they typically have it in PATH

    def extract_text_from_docx(self, docx_path):
        """Extract text from DOCX files."""
        try:
            doc = docx.Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            return ""

    def extract_text_from_image(self, image_path):
        """Extract text from images using OCR."""
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {str(e)}")
            return ""

    def process_document(self, file_path):
        """Process document and extract text based on file type."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""

        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff']:
                text = self.extract_text_from_image(file_path)
            else:
                # Try to read as plain text for other file types
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text = file.read()
                except Exception as e:
                    logger.error(f"Could not read file as text: {str(e)}")
                    return ""

            # Clean and normalize the extracted text
            text = self.clean_text(text)
            
            # Log the extracted text length for debugging
            logger.info(f"Extracted {len(text)} characters from {file_path}")
            
            return text

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return ""

    def clean_text(self, text):
        """Clean and normalize extracted text with improved processing."""
        if not text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Replace common OCR mistakes
        ocr_fixes = {
            '0': 'o',
            '1': 'i',
            '5': 's',
            '@': 'a',
            '&': 'and',
        }
        
        for wrong, right in ocr_fixes.items():
            text = text.replace(wrong, right)
        
        # Remove special characters but keep spaces and basic punctuation
        text = re.sub(r'[^a-z0-9\s.,!?-]', ' ', text)
        
        # Normalize whitespace
        text = ' '.join(text.split())
        
        # Log cleaned text length
        logger.info(f"Cleaned text length: {len(text)} characters")
        
        return text

    def detect_clearance_type(self, text):
        """Detect clearance type based on keywords with improved logging."""
        matches = {}
        text_lower = text.lower()
        
        # Log the first 200 characters of processed text for debugging
        logger.info(f"Analyzing text (first 200 chars): {text_lower[:200]}...")
        
        for clearance_type, keywords in self.clearance_keywords.items():
            count = 0
            found_keywords = []
            
            for keyword in keywords:
                if keyword.lower() in text_lower:
                    count += 1
                    found_keywords.append(keyword)
            
            matches[clearance_type] = {
                'count': count,
                'found_keywords': found_keywords
            }
            
            # Log matches for each type
            if found_keywords:
                logger.info(f"{clearance_type} matches: {count} keywords found: {', '.join(found_keywords)}")

        # Find the type with the most matches
        if any(m['count'] for m in matches.values()):
            best_match = max(matches.items(), key=lambda x: x[1]['count'])
            if best_match[1]['count'] >= 2:  # Require at least 2 keyword matches
                logger.info(f"Selected clearance type: {best_match[0]} with {best_match[1]['count']} matches")
                return best_match[0]
            else:
                logger.info("No clearance type had enough keyword matches (minimum 2 required)")
        else:
            logger.info("No keywords matched in the document")
        
        return None

    def analyze_document(self, file_path):
        """Analyze document with improved logging."""
        text = self.process_document(file_path)
        
        if not text:
            logger.error(f"No text extracted from document: {file_path}")
            return {
                'status': 'ERROR',
                'message': 'Could not extract text from document',
                'clearance_type': None,
                'text_content': None
            }

        # Log the text content for debugging (first 500 characters)
        logger.info(f"Extracted text sample: {text[:500]}...")
        
        clearance_type = self.detect_clearance_type(text)
        
        result = {
            'status': 'SUCCESS',
            'message': 'Document processed successfully',
            'clearance_type': clearance_type,
            'text_content': text
        }
        
        logger.info(f"Analysis result: {result['status']}, Clearance Type: {result['clearance_type']}")
        
        return result

    def extract_academic_sections(self, text):
        """Extract and analyze specific sections from academic documents."""
        sections = {
            'academic_affairs': '',
            'library': '',
            'department': ''
        }
        
        text_lower = text.lower()
        
        # Try to extract academic affairs section
        if 'academic affairs' in text_lower:
            start = text_lower.find('academic affairs')
            end = text_lower.find('library', start)
            if end == -1:
                end = len(text_lower)
            sections['academic_affairs'] = text[start:end].strip()
        
        # Try to extract library section
        if 'library' in text_lower:
            start = text_lower.find('library')
            end = text_lower.find('department', start)
            if end == -1:
                end = len(text_lower)
            sections['library'] = text[start:end].strip()
        
        # Try to extract department section
        if 'department' in text_lower:
            start = text_lower.find('department')
            end = text_lower.find('date', start)
            if end == -1:
                end = len(text_lower)
            sections['department'] = text[start:end].strip()
        
        return sections

class ClearancePredictor:
    def __init__(self):
        self.document_processor = DocumentProcessor()

    def predict(self, document):
        try:
            logger.info(f"Starting prediction for document {document.id}")
            
            if not document.file:
                logger.error(f"No file attached to document {document.id}")
                return 'PENDING'

            # Process the document
            analysis = self.document_processor.analyze_document(document.file.path)
            
            if analysis['status'] == 'ERROR':
                logger.error(f"Error processing document {document.id}: {analysis['message']}")
                return 'PENDING'

            text_lower = analysis['text_content'].lower()
            
            # Enhanced approval indicators with more variations
            approval_indicators = [
                'approved', 'cleared', 'completed', 'verified',
                'confirmed', 'accepted', 'passed', 'satisfactory',
                'requirements met', 'obligations fulfilled', 'granted',
                'authorized', 'endorsed', 'certified', 'validated',
                'received and reviewed',  # Added based on your document content
                'satisfied', 'complied with',  # Added based on your document content
                'submitted', 'processed',
                'i have received', 'i am satisfied',  # Added based on your document content
                'this is to confirm'  # Added based on your document content
            ]

            # Enhanced rejection indicators
            rejection_indicators = [
                'rejected', 'denied', 'failed', 'incomplete',
                'not approved', 'pending requirements', 'outstanding',
                'missing', 'unsatisfactory', 'declined',
                'not received', 'not satisfied', 'not complied',
                'not submitted', 'pending submission'
            ]

            # Count matches with context
            approval_matches = []
            rejection_matches = []
            
            # Check for approval phrases
            for indicator in approval_indicators:
                if indicator in text_lower:
                    approval_matches.append(indicator)
                    
            # Check for rejection phrases
            for indicator in rejection_indicators:
                if indicator in text_lower:
                    rejection_matches.append(indicator)

            # Log matched phrases
            if approval_matches:
                logger.info(f"Found approval phrases: {', '.join(approval_matches)}")
            if rejection_matches:
                logger.info(f"Found rejection phrases: {', '.join(rejection_matches)}")

            # Enhanced decision logic
            if len(approval_matches) > 0 and len(rejection_matches) == 0:
                # If we have approval indicators and no rejection indicators
                status = 'APPROVED'
            elif len(rejection_matches) > 0:
                # If we have any rejection indicators
                status = 'REJECTED'
            elif 'this is to confirm' in text_lower and 'received' in text_lower:
                # Special case for confirmation statements
                status = 'APPROVED'
            elif analysis['clearance_type'] is not None:
                # If we detected a clearance type but no clear approval/rejection
                # Check for specific clearance type indicators
                if analysis['clearance_type'] == 'ACADEMIC' and any(phrase in text_lower for phrase in [
                    'course file', 'received and reviewed', 'satisfied'
                ]):
                    status = 'APPROVED'
                elif analysis['clearance_type'] == 'LIBRARY' and any(phrase in text_lower for phrase in [
                    'books returned', 'no outstanding'
                ]):
                    status = 'APPROVED'
                else:
                    status = 'PENDING'
            else:
                status = 'PENDING'

            logger.info(f"Document {document.id} - Found {len(approval_matches)} approval indicators and {len(rejection_matches)} rejection indicators")
            logger.info(f"Final prediction for document {document.id}: {status}")
            return status

        except Exception as e:
            logger.error(f"Error predicting document {document.id}: {str(e)}")
            return 'PENDING'
